"""
Parses a Golden Rules playbook .docx (e.g. "Freo Group AU - Review Playbook
(AI-ready).docx") into structured JSON for the contract-drafting tool
("Job B") and, eventually, a contract-review engine ("Job A").

Verified live against the Freo Group AU document: every rule is a paragraph
pair immediately followed by one 6-row table, in that exact order in the
document body —
    [P] "{RULE_ID}   {Title}"
    [P] "Priority: {P}     Applies to: {A}"
    [TABLE] WHERE TO LOOK | REQUIRED | FALLBACK | ESCALATE IF | FLAG IF | PREFERRED LANGUAGE
Category is derived from the rule-ID prefix (confirmed 1:1 with the
document's own category headings) rather than parsed from prose, since the
prefix mapping is far more robust than trying to track "which heading came
last" while walking the body.

Usage:
    python playbook_parser.py "<path to playbook docx>" --id freo-group-au --label "Freo Group AU - Crane Hire"
"""

import argparse
import json
import re
from pathlib import Path

import docx
from docx.oxml.ns import qn

RULE_HEADER_RE = re.compile(r"^([A-Z]{2,4}-\d{2})\s+(.+)$")
PRIORITY_APPLIES_RE = re.compile(r"^Priority:\s*(.+?)\s+Applies to:\s*(.+)$")
SOURCE_LINE_RE = re.compile(r"^source:\s*", re.IGNORECASE)

CATEGORY_BY_PREFIX = {
    "LIA": "Liability & Indemnity",
    "LD": "Liquidated Damages & Delay",
    "DEF": "Defects & Warranties",
    "SEC": "Security & Guarantees",
    "PAY": "Payment & Money",
    "TRM": "Standby, Suspension & Termination",
    "OPS": "Scope, Site & Operations",
    "INS": "Insurance",
    "EQP": "Equipment",
    "IPC": "IP, Confidentiality & Data",
    "CMR": "Commercial & Structural",
    "FLW": "Flow-down & Upstream Risk",
}

TABLE_ROW_ORDER = ["WHERE TO LOOK", "REQUIRED", "FALLBACK", "ESCALATE IF", "FLAG IF", "PREFERRED LANGUAGE"]


def _category_for(rule_id: str) -> str:
    prefix = rule_id.split("-")[0]
    if prefix not in CATEGORY_BY_PREFIX:
        raise ValueError(f"Unknown rule-ID prefix {prefix!r} for rule {rule_id!r} — "
                          f"CATEGORY_BY_PREFIX needs a new entry.")
    return CATEGORY_BY_PREFIX[prefix]


def _paragraph_text(p_element) -> str:
    return "".join(node.text or "" for node in p_element.iter(qn("w:t")))


def _parse_preferred_language(cell) -> tuple[str | None, str | None]:
    paras = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    if not paras:
        return None, None
    if SOURCE_LINE_RE.match(paras[-1]):
        source_tag = SOURCE_LINE_RE.sub("", paras[-1]).strip()
        language = " ".join(paras[:-1]).strip()
        return (language or None), source_tag
    # No trailing "Source:" line means this rule has no insertable model
    # language at all (verified: the document's own convention for its 30
    # "flag and describe, not draft" rules) — the explanatory text present
    # (e.g. "None - amend or delete as described above") isn't clause
    # content and shouldn't be treated as any.
    return None, None


def _parse_rule_table(table) -> dict:
    by_label = {row.cells[0].text.strip().upper(): row.cells[1] for row in table.rows}
    missing = [label for label in TABLE_ROW_ORDER if label not in by_label]
    if missing:
        raise ValueError(f"Rule table missing expected row(s): {missing}")

    flag_if = [p.text.strip() for p in by_label["FLAG IF"].paragraphs if p.text.strip()]
    preferred_language, source_tag = _parse_preferred_language(by_label["PREFERRED LANGUAGE"])

    return {
        "where_to_look": by_label["WHERE TO LOOK"].text.strip(),
        "required": by_label["REQUIRED"].text.strip(),
        "fallback": by_label["FALLBACK"].text.strip(),
        "escalate_if": by_label["ESCALATE IF"].text.strip(),
        "flag_if": flag_if,
        "preferred_language": preferred_language,
        "source_tag": source_tag,
    }


def parse_playbook(docx_path: str) -> list[dict]:
    doc = docx.Document(docx_path)
    table_iter = iter(doc.tables)

    rules = []
    pending = None
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = _paragraph_text(child).strip()
            if not text:
                continue
            header_match = RULE_HEADER_RE.match(text)
            if header_match and header_match.group(1).split("-")[0] in CATEGORY_BY_PREFIX:
                pending = {"rule_id": header_match.group(1), "title": header_match.group(2).strip()}
                continue
            priority_match = PRIORITY_APPLIES_RE.match(text)
            if priority_match and pending is not None:
                pending["priority"] = priority_match.group(1).strip()
                pending["applies_to"] = priority_match.group(2).strip()
        elif child.tag == qn("w:tbl"):
            table = next(table_iter)
            if pending is not None and "priority" in pending:
                rule = {**pending, "category": _category_for(pending["rule_id"])}
                rule.update(_parse_rule_table(table))
                rules.append(rule)
                pending = None
            # Tables not immediately preceded by a completed rule header
            # (e.g. the source-tag summary table at the top of the doc) are
            # reference material, not rules — skip them.

    return rules


def _assert_expected_shape(rules: list[dict]) -> None:
    """Regression guard: fails loudly if a future edit to the source .docx
    changes the document's structure in a way this parser doesn't handle,
    rather than silently producing a partial/wrong rule set."""
    assert len(rules) == 86, f"expected 86 rules, parsed {len(rules)}"
    ids = [r["rule_id"] for r in rules]
    assert len(ids) == len(set(ids)), "duplicate rule_id detected"

    applies_to_counts = {}
    priority_counts = {}
    for r in rules:
        applies_to_counts[r["applies_to"]] = applies_to_counts.get(r["applies_to"], 0) + 1
        priority_counts[r["priority"]] = priority_counts.get(r["priority"], 0) + 1
    assert applies_to_counts == {
        "All contract types": 57, "Wind / renewables subcontract": 23,
        "Mining master supply agreement": 3, "Equipment hire": 3,
    }, f"applies_to distribution changed: {applies_to_counts}"
    assert priority_counts == {
        "PRESS": 44, "MUST PRESS": 36, "MANAGE": 5, "ACCEPT+NOTE": 1,
    }, f"priority distribution changed: {priority_counts}"

    with_language = [r for r in rules if r["preferred_language"]]
    assert len(with_language) == 56, f"expected 56 rules with preferred_language, got {len(with_language)}"
    source_tag_counts = {}
    for r in with_language:
        source_tag_counts[r["source_tag"]] = source_tag_counts.get(r["source_tag"], 0) + 1
    assert source_tag_counts == {
        "External counsel": 28, "Unvetted draft - counsel review needed": 12,
        "Freo register": 8, "Executed contract": 8,
    }, f"source_tag distribution changed: {source_tag_counts}"

    bracket_re = re.compile(r"\[[^\]]+\]")
    bracket_rules = [r["rule_id"] for r in with_language if bracket_re.search(r["preferred_language"])]
    assert len(bracket_rules) == 6, f"expected 6 rules with bracket placeholders, got {len(bracket_rules)}: {bracket_rules}"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path", help="Path to the playbook .docx")
    parser.add_argument("--id", required=True, help="Short id for this playbook, e.g. freo-group-au")
    parser.add_argument("--label", required=True, help="Display label, e.g. 'Freo Group AU - Crane Hire'")
    parser.add_argument("--skip-shape-check", action="store_true",
                         help="Skip the Freo-Group-AU-specific count assertions (use for a different playbook doc)")
    args = parser.parse_args()

    rules = parse_playbook(args.docx_path)
    print(f"Parsed {len(rules)} rules from {args.docx_path}")

    if not args.skip_shape_check:
        _assert_expected_shape(rules)
        print("Shape check passed (86 rules, expected applies_to/priority/source_tag distributions).")

    out_dir = Path(__file__).parent.parent / "mclegal-frontend" / "public" / "playbooks"
    out_dir.mkdir(parents=True, exist_ok=True)

    rules_path = out_dir / f"{args.id}.json"
    rules_path.write_text(json.dumps(rules, indent=2), encoding="utf-8")
    print(f"Wrote {rules_path}")

    manifest_path = out_dir / "manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8")) if manifest_path.exists() else []
    manifest = [m for m in manifest if m["id"] != args.id]
    contract_types = sorted({r["applies_to"] for r in rules if r["applies_to"] != "All contract types"})
    manifest.append({"id": args.id, "label": args.label, "contractTypes": contract_types, "file": f"{args.id}.json"})
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(f"Wrote {manifest_path}")


if __name__ == "__main__":
    main()
