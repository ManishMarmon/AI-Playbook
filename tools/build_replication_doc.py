"""
Builds McLegal_Replication_Guide.docx — the complete replication
specification for this project.

The document has two halves:
  1. Narrative (architecture, runbooks, config contracts, design decisions)
     supplied as structured JSON (see --narrative), drafted by subsystem
     research and rendered into headings/paragraphs/tables.
  2. Mechanical byte-exact embedding of every source file: each file is
     written between marker paragraphs with its SHA-256, EOL style, and
     trailing-newline flag recorded, so tools/extract_from_docx.py can
     re-materialize the whole tree on a new machine and PROVE it identical.

Files that can't round-trip as text (binaries, non-UTF-8, mixed EOLs) are
embedded as base64. Secrets are never embedded: .env is excluded by
design; only the placeholder .env.example files are included.

Usage (from repo root):
    python tools/build_replication_doc.py --narrative <sections.json> --out McLegal_Replication_Guide.docx
"""

import argparse
import base64
import hashlib
import json
import re
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
BEGIN = "⟦BEGIN-FILE⟧"
END = "⟦END-FILE⟧"
W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

# ── Embed manifest ───────────────────────────────────────────────────────────
# Globs are resolved at build time so newly added files are never silently
# missed. .env (real secrets) is deliberately absent.

TEXT_GROUPS = [
    ("Root documentation and configuration", [
        "README.md", "CONTEXT.md", "AZURE_OPENAI_PORT_PLAN.md", "LLM_COST_LOG.md",
        ".env.example", ".gitignore", ".github/workflows/ci.yml",
    ]),
    ("Local-only project records (gitignored — this document is their only transfer path)", [
        # docs/audit_findings_2026-08-25.json is deliberately NOT embedded: it
        # quotes the real Key Vault URL / tenant hostname it audited (that is
        # its job as a record). It ships via the copy-from-original manifest.
        "docs/AUDIT_FINDINGS.md", "docs/AUDIT_FINDINGS_2026-08-25.md",
        "docs/REMEDIATION_PLAN.md",
    ]),
    ("Tools", ["tools/*.py"]),
    ("Backend — redline_discovery", [
        "redline_discovery/*.py", "redline_discovery/requirements.txt",
        "redline_discovery/constraints.txt",
        "redline_discovery/pytest.ini", "redline_discovery/db/*.sql",
    ]),
    ("Backend tests", ["redline_discovery/tests/*.py"]),
    ("Claude Code workflow scripts", ["redline_discovery/workflows/*.js"]),
    ("Frontend — mclegal-frontend configuration", [
        "mclegal-frontend/package.json", "mclegal-frontend/package-lock.json",
        "mclegal-frontend/tsconfig.json", "mclegal-frontend/vite.config.ts",
        "mclegal-frontend/index.html", "mclegal-frontend/.env.example",
    ]),
    ("Frontend — source", ["mclegal-frontend/src/**/*"]),
]

PRODUCT_GROUPS = [
    ("Playbook product artifacts (LLM-derived — use these copies, do not regenerate)", [
        "mclegal-frontend/public/playbooks/*.json",
    ]),
    ("Small generated data files", [
        "mclegal-frontend/public/data/analytics.json",
        "mclegal-frontend/public/data/discovery_summary.json",
    ]),
]

BINARY_FILES = [
    "mclegal-frontend/public/marmon-mark-white.png",
    "mclegal-frontend/public/marmon-mark-blue.png",
    "small_logo1.png", "small_logo2.jpg", "marmon_logo.png",
    # Gitignored source documents a git clone would MISS — the Freo rulebook
    # is the B2 golden-rules schema template playbook_parser.py parses.
    "Freo Group AU - Review Playbook (AI-ready).docx",
    "Playbook_AI_Project_Redline_Document_Identification.docx",
]

# Large derived data files: NOT embedded; the doc explains copy/regenerate.
LARGE_DERIVED = [
    "mclegal-frontend/public/data/requests_catalog.json",
    "mclegal-frontend/public/data/clause_findings.json",
    "mclegal-frontend/public/data/golden_rules_findings.json",
    "mclegal-frontend/public/data/redline_catalog.json",
    "mclegal-frontend/public/data/redline_diffs.json",
]


def resolve_group(patterns):
    seen = []
    for pat in patterns:
        matches = sorted(p for p in ROOT.glob(pat) if p.is_file()) if any(c in pat for c in "*?[") else [ROOT / pat]
        for p in matches:
            if p.is_file() and "__pycache__" not in p.parts and p.suffix != ".pyc":
                rel = p.relative_to(ROOT).as_posix()
                if rel not in seen:
                    seen.append(rel)
    return seen


# ── Byte-exact classification ────────────────────────────────────────────────

_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def classify(content: bytes) -> dict:
    sha = hashlib.sha256(content).hexdigest()
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        return {"mode": "b64", "sha": sha, "lines": _b64_lines(content), "eol": "NA", "eofnl": "NA"}
    if _CTRL.search(text):
        return {"mode": "b64", "sha": sha, "lines": _b64_lines(content), "eol": "NA", "eofnl": "NA"}
    has_crlf = "\r\n" in text
    lone_lf = re.search(r"(?<!\r)\n", text) is not None
    lone_cr = re.search(r"\r(?!\n)", text) is not None
    if lone_cr or (has_crlf and lone_lf):
        return {"mode": "b64", "sha": sha, "lines": _b64_lines(content), "eol": "NA", "eofnl": "NA"}
    eol = "CRLF" if has_crlf else "LF"
    sep = "\r\n" if has_crlf else "\n"
    eofnl = "1" if text.endswith(sep) else "0"
    if eofnl == "1":
        text = text[: -len(sep)]
    lines = text.split(sep) if text or eofnl == "1" else [""]
    return {"mode": "text", "sha": sha, "lines": lines, "eol": eol, "eofnl": eofnl}


def _b64_lines(content: bytes, width: int = 96):
    b = base64.b64encode(content).decode("ascii")
    return [b[i:i + width] for i in range(0, len(b), width)] or [""]


# ── Low-level docx helpers ───────────────────────────────────────────────────

def add_raw_paragraph(body_sectpr, text: str, style_id: str):
    if text:
        inner = []
        for i, part in enumerate(text.split("\t")):
            if i:
                inner.append("<w:tab/>")
            if part:
                inner.append(f'<w:t xml:space="preserve">{escape(part)}</w:t>')
        runs = "<w:r>" + "".join(inner) + "</w:r>"
    else:
        runs = ""
    xml = f'<w:p {W_NS}><w:pPr><w:pStyle w:val="{style_id}"/></w:pPr>{runs}</w:p>'
    body_sectpr.addprevious(parse_xml(xml))


def make_styles(doc):
    def para_style(name, font, size, bold=False, color=None, before=0, after=0):
        s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        s.font.name, s.font.size, s.font.bold = font, Pt(size), bold
        if color:
            s.font.color.rgb = RGBColor(*color)
        pf = s.paragraph_format
        pf.space_before, pf.space_after, pf.line_spacing = Pt(before), Pt(after), 1.0
        return s
    para_style("CodeLine", "Consolas", 7)
    para_style("FileMarker", "Consolas", 8, bold=True, color=(0x1F, 0x44, 0xC4), before=10, after=2)
    para_style("CmdBlock", "Consolas", 9, before=4, after=4)


def add_toc(doc):
    xml = (
        f'<w:p {W_NS}><w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:t>Table of contents: select this line, then press F9 (or right-click, Update Field) in Word.</w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
    )
    doc.element.body.xpath("./w:sectPr")[0].addprevious(parse_xml(xml))


def embed_file(doc, rel: str, manifest: list):
    content = (ROOT / rel).read_bytes()
    info = classify(content)
    manifest.append((rel, info["sha"], len(content), info["mode"]))
    sectpr = doc.element.body.xpath("./w:sectPr")[0]
    marker = (f"{BEGIN} path={rel} mode={info['mode']} eol={info['eol']} "
              f"eofnl={info['eofnl']} sha256={info['sha']}")
    add_raw_paragraph(sectpr, marker, "FileMarker")
    for line in info["lines"]:
        add_raw_paragraph(sectpr, line, "CodeLine")
    add_raw_paragraph(sectpr, END, "FileMarker")


# ── Narrative rendering ──────────────────────────────────────────────────────

def render_blocks(doc, blocks):
    sectpr = doc.element.body.xpath("./w:sectPr")[0]
    for b in blocks or []:
        kind = b.get("type")
        if kind == "para" and b.get("text"):
            doc.add_paragraph(b["text"])
        elif kind == "bullets":
            for item in b.get("items") or []:
                doc.add_paragraph(item, style="List Bullet")
        elif kind == "code" and b.get("text") is not None:
            for line in b["text"].split("\n"):
                add_raw_paragraph(sectpr, line, "CmdBlock")
        elif kind == "table":
            header = b.get("header") or []
            rows = b.get("rows") or []
            if not header and rows:
                header = [""] * len(rows[0])
            if not header:
                continue
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Table Grid"
            for i, h in enumerate(header):
                run = t.rows[0].cells[i].paragraphs[0].add_run(str(h))
                run.bold = True
                run.font.size = Pt(8.5)
            for row in rows:
                cells = t.add_row().cells
                for i in range(len(header)):
                    val = str(row[i]) if i < len(row) else ""
                    r = cells[i].paragraphs[0].add_run(val)
                    r.font.size = Pt(8.5)
            doc.add_paragraph()


def render_sections(doc, sections):
    for sec in sections or []:
        doc.add_heading(sec.get("heading", ""), level=max(2, min(4, int(sec.get("level", 2)))))
        render_blocks(doc, sec.get("blocks"))


# ── Document assembly ────────────────────────────────────────────────────────

def build(narrative: dict, out_path: Path):
    doc = Document()
    make_styles(doc)
    core = doc.core_properties
    core.title = "McLegal / Redline Discovery — Replication Guide"
    core.subject = "Complete environment + source + data specification for identical rebuild"

    doc.add_heading("McLegal / Redline Discovery Engine", level=0)
    doc.add_paragraph("Complete Replication Guide — environment, database, pipeline, frontend, and byte-exact source embedding.")
    doc.add_paragraph("CONFIDENTIAL — Marmon internal. Contains full application source and internal design detail. "
                      "Secrets and internal hostnames are deliberately NOT in this document: every such value is a "
                      "placeholder to be filled from the original laptop's gitignored .env files or Azure Key Vault.")
    add_toc(doc)
    doc.add_page_break()

    # Part 0 — how to use
    doc.add_heading("Part 0 — How to Use This Document", level=1)
    p0 = narrative.get("part0", {})
    render_sections(doc, p0.get("sections"))

    order = [
        ("Part 1 — Project Overview and Design Decisions", "ov"),
        ("Part 2 — Prerequisites and Environment", "env"),
        ("Part 3 — Backend Data Layer, Configuration, and Database", "data"),
        ("Part 4 — Deterministic Analysis Pipeline", "det"),
        ("Part 5 — LLM Pipeline and Playbook Production", "llm"),
        ("Part 6 — Frontend Application", "fe"),
        ("Part 7 — Master Rebuild Checklist and Verification", "checklist"),
    ]
    for title, key in order:
        doc.add_page_break()
        doc.add_heading(title, level=1)
        render_sections(doc, (narrative.get(key) or {}).get("sections"))

    # Part 8 — manifest
    doc.add_page_break()
    doc.add_heading("Part 8 — Embedded File Manifest and Integrity", level=1)
    doc.add_paragraph("Every file below is embedded byte-exactly in the appendices between "
                      f"'{BEGIN}' / '{END}' marker paragraphs. Run tools/extract_from_docx.py against this .docx "
                      "to re-materialize and hash-verify all of them. The manifest is repeated here so a human can "
                      "audit coverage at a glance (sha256 · bytes · mode · path).")
    manifest: list = []
    manifest_anchor = len(doc.element.body)  # filled after embedding; see below

    # Appendix A/B — embed everything
    doc.add_page_break()
    doc.add_heading("Appendix A — Complete Source Code (byte-exact)", level=1)
    for group, patterns in TEXT_GROUPS:
        doc.add_heading(group, level=2)
        for rel in resolve_group(patterns):
            embed_file(doc, rel, manifest)

    doc.add_page_break()
    doc.add_heading("Appendix B — Product Data Artifacts (byte-exact)", level=1)
    doc.add_paragraph("The playbook JSONs are LLM-derived product artifacts reviewed in context — re-running the "
                      "pipeline produces semantically similar but NOT identical rules. Use these embedded copies.")
    for group, patterns in PRODUCT_GROUPS:
        doc.add_heading(group, level=2)
        for rel in resolve_group(patterns):
            embed_file(doc, rel, manifest)

    doc.add_page_break()
    doc.add_heading("Appendix C — Binary Assets (base64, byte-exact)", level=1)
    doc.add_paragraph("Rendered previews first (for the human reader), then base64 blocks (for the extractor). "
                      "The extractor decodes these back to identical binaries.")
    for rel in BINARY_FILES:
        if (ROOT / rel).exists():
            doc.add_paragraph(rel + ":")
            try:
                doc.add_picture(str(ROOT / rel), width=Inches(1.6))
            except Exception:
                doc.add_paragraph("(preview unavailable)")
    for rel in BINARY_FILES:
        if (ROOT / rel).exists():
            embed_file(doc, rel, manifest)

    # Now back-fill the Part 8 manifest listing (append at anchor position)
    sectpr = doc.element.body.xpath("./w:sectPr")[0]
    anchor_children = list(doc.element.body)
    insert_before = anchor_children[manifest_anchor] if manifest_anchor < len(anchor_children) else sectpr
    for rel, sha, size, mode in manifest:
        xml_line = f"{sha}  {size:>9,}  {mode:<4}  {rel}"
        p = parse_xml(f'<w:p {W_NS}><w:pPr><w:pStyle w:val="CodeLine"/></w:pPr>'
                      f'<w:r><w:t xml:space="preserve">{escape(xml_line)}</w:t></w:r></w:p>')
        insert_before.addprevious(p)

    doc.save(out_path)
    return manifest


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--narrative", default=None, help="Path to sections JSON (keys: part0, ov, env, data, det, llm, fe, checklist)")
    ap.add_argument("--out", default=str(ROOT / "McLegal_Replication_Guide.docx"))
    args = ap.parse_args()

    narrative = {}
    if args.narrative:
        narrative = json.loads(Path(args.narrative).read_text(encoding="utf-8"))
    mf = build(narrative, Path(args.out))
    total = sum(m[2] for m in mf)
    print(f"Built {args.out}\n{len(mf)} files embedded, {total:,} source bytes")


if __name__ == "__main__":
    main()
